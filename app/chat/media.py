from __future__ import annotations

import base64
from io import BytesIO
from typing import Any

from fastapi import UploadFile
from openai import AsyncOpenAI

from app.core.config import get_settings


MAX_EXTRACTED_CHARS = 30_000

_client: AsyncOpenAI | None = None


async def media_to_part(media: UploadFile) -> dict[str, Any]:
    mime = media.content_type or ""
    data = await media.read()

    if mime.startswith("image/"):
        b64 = base64.b64encode(data).decode()
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        }

    if mime.startswith("audio/") or mime == "application/ogg":
        transcript = await whisper_transcribe(data, media.filename or "audio.ogg")
        return {"type": "text", "text": f"[пользователь сказал голосом]:\n{transcript}"}

    if mime == "application/pdf":
        return {
            "type": "text",
            "text": f"[документ PDF]:\n{extract_pdf_text(data)[:MAX_EXTRACTED_CHARS]}",
        }

    if mime.endswith("wordprocessingml.document"):
        return {
            "type": "text",
            "text": f"[документ DOCX]:\n{extract_docx_text(data)[:MAX_EXTRACTED_CHARS]}",
        }

    raise ValueError(f"Unsupported media type: {mime}")


async def whisper_transcribe(audio_bytes: bytes, filename: str) -> str:
    f = BytesIO(audio_bytes)
    f.name = filename
    result = await _get_client().audio.transcriptions.create(model="whisper-1", file=f)
    return result.text


def extract_pdf_text(data: bytes, max_pages: int = 50) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = reader.pages[:max_pages]
    text_parts = [(page.extract_text() or "").strip() for page in pages]
    text = "\n\n".join(part for part in text_parts if part)

    if len(pages) >= 5 and len(text) < 100:
        return "[похоже, PDF является сканом: извлечено меньше 100 символов текста]"
    return text


def extract_docx_text(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _get_client() -> AsyncOpenAI:
    global _client
    if _client is None:
        settings = get_settings()
        _client = AsyncOpenAI(
            api_key=settings.openai_api_key.get_secret_value(),
            base_url=settings.openai_base_url,
            timeout=settings.request_timeout,
        )
    return _client
